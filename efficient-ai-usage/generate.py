import subprocess
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import textwrap

OUT_DIR = Path(__file__).parent / "build"
OUT_DIR.mkdir(exist_ok=True)
PLANTUML_JAR = r"C:\tools\plantuml.jar"

# ---------------------------------------------------------------------------
# PlantUML sources
# ---------------------------------------------------------------------------

DIAGRAMS = {
    "agent_pipeline": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
skinparam BoxPadding 12

title Multi-Agent Pipeline — Model per Step

rectangle "Incoming Task" as IN #E8F4FD

rectangle "Agent 1\\nLocal Model\\n(Phi-3 / Mistral 7B)\\nClassify + Pre-process" as A1 #D5E8D4
rectangle "Agent 2\\nEmbedding Model\\n+ Vector DB\\nRetrieve top-K chunks" as A2 #D5E8D4
rectangle "Agent 3\\nMid-tier Model\\n(Sonnet / Flash)\\nIntermediate reasoning" as A3 #FFE6CC
rectangle "Agent 4\\nPremium Model\\n(Opus / Ultra)\\nSynthesize & Generate" as A4 #F8CECC
rectangle "Agent 5\\nSmall Model\\n(Haiku / Flash)\\nFormat & Validate schema" as A5 #D5E8D4

rectangle "Structured Output" as OUT #E8F4FD

IN --> A1
A1 --> A2
A2 --> A3
A3 --> A4 : escalate only\\nif needed
A4 --> A5
A5 --> OUT

note right of A4
  Premium model fires
  only at this step
end note
@enduml
""",

    "rag_flow": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
skinparam sequenceMessageAlign center

title RAG — Ground the Model in What It Cannot Know

participant "User Query" as Q
participant "Embedding\\nModel" as EM
participant "Vector DB" as VDB
participant "Re-ranker" as RR
participant "Premium\\nModel" as LLM
participant "Response" as R

Q -> EM : embed query
EM -> VDB : similarity search
VDB --> RR : top-20 candidates
RR --> LLM : top-K chunks (3-5)\\nonly relevant context
LLM --> R : grounded answer

note over VDB
  Stores: SDK docs, kernel
  internals, internal design
  docs, version-pinned refs
end note

note over LLM
  Reasons from injected
  context, not latent space
end note
@enduml
""",

    "prompt_caching": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title Prompt Caching — Cost Per Query

rectangle "Without Caching" as WOC {
  rectangle "System Prompt\\n800 tokens" as SP1 #F8CECC
  rectangle "KB Chunk\\n12,000 tokens" as KB1 #F8CECC
  rectangle "User Question\\n50 tokens" as UQ1 #F8CECC
}

rectangle "With Caching" as WC {
  rectangle "System Prompt\\n800 tokens\\n[CACHED — 10% cost]" as SP2 #D5E8D4
  rectangle "KB Chunk\\n12,000 tokens\\n[CACHED — 10% cost]" as KB2 #D5E8D4
  rectangle "User Question\\n50 tokens\\n[billed normally]" as UQ2 #FFE6CC
}

note bottom of WOC
  Every query: 12,850 tokens billed
end note

note bottom of WC
  Every query after first: ~50 tokens + cache read fee
  Saving: >80% on input costs
end note
@enduml
""",

    "tool_calling": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
skinparam sequenceMessageAlign center

title Programmable Tool Calling — Model as Orchestrator

participant "Premium\\nModel" as LLM
participant "Tool:\\nvector_search()" as T1 #D5E8D4
participant "Tool:\\nclassify_intent()" as T2 #D5E8D4
participant "Tool:\\nrun_local_model()" as T3 #D5E8D4
participant "Tool:\\nfetch_doc_chunk()" as T4 #D5E8D4
participant "Tool:\\nformat_output()" as T5 #D5E8D4

LLM -> T1 : call (parallel)
LLM -> T2 : call (parallel)
T1 --> LLM : top-K chunks
T2 --> LLM : intent + confidence

alt confidence HIGH — routine task
  LLM -> T3 : delegate to local model
  T3 --> LLM : structured result
else confidence LOW — complex task
  LLM -> T4 : fetch additional context
  T4 --> LLM : doc chunk
  LLM -> LLM : reason over full context
end

LLM -> T5 : enforce output schema
T5 --> LLM : validated JSON

note over T1, T2
  Parallel tool calls = one
  round trip, not two
end note

note over T3
  Local model invoked as a tool —
  premium model never sees
  the raw processing cost
end note
@enduml
""",

    "full_stack": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
skinparam rectangleBorderColor #888888

title The Full Efficient AI Stack

rectangle "Incoming Task" as T0 #E8F4FD

rectangle "Step 1: Classify + Pre-process\\nLocal Model — free, fast" as T1 #D5E8D4
rectangle "Step 2: Vector DB Retrieval\\nEmbedding model + top-K only" as T2 #D5E8D4
rectangle "Step 3: Cache Check\\nIs prompt prefix warm?" as T3 #FFF2CC
rectangle "Step 4: Tool Calling Layer\\nParallel calls: retrieve + classify" as T4 #FFF2CC
rectangle "Step 5: Premium Model\\nSynthesize — only what requires it" as T5 #F8CECC
rectangle "Step 6: Format + Validate\\nTool schema enforcement" as T6 #D5E8D4
rectangle "Structured Output" as T7 #E8F4FD

T0 --> T1
T1 --> T2
T2 --> T3
T3 --> T4
T4 --> T5 : escalate if\\ncomplexity high
T5 --> T6
T6 --> T7

note right of T3
  Cache hit: pay 10%
  Cache miss: re-warm prefix
end note

note right of T5
  Async Batching API:
  50% off for non-realtime
  workloads
end note
@enduml
""",
}


def render_diagrams():
    print("Rendering PlantUML diagrams...")
    paths = {}
    for name, src in DIAGRAMS.items():
        puml_path = OUT_DIR / f"{name}.puml"
        png_path = OUT_DIR / f"{name}.png"
        puml_path.write_text(src.strip(), encoding="utf-8")
        result = subprocess.run(
            ["java", "-jar", PLANTUML_JAR, str(puml_path), "-o", str(OUT_DIR)],
            capture_output=True, text=True
        )
        if result.returncode != 0:
            print(f"  ERROR rendering {name}: {result.stderr}")
        else:
            print(f"  OK: {name}.png")
        paths[name] = png_path
    return paths


# ---------------------------------------------------------------------------
# Word document helpers
# ---------------------------------------------------------------------------

def set_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F0F0F0")
    p._p.get_or_add_pPr().append(shading)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_image(doc, path, width_inches=5.5, caption=None):
    if not Path(path).exists():
        doc.add_paragraph(f"[Diagram not rendered: {path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(9)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "D5E8D4")
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build_doc(diagram_paths):
    doc = Document()

    # Title
    title = doc.add_heading("Stop Burning Your AI Budget", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("An Engineer's Guide to Efficient Premium Model Usage")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    sub.runs[0].font.size = Pt(13)
    doc.add_paragraph()

    # Intro
    add_body(doc,
        "The debate is everywhere in enterprise right now: are we getting productivity returns "
        "that justify the token costs? Finance wants invoices justified. Engineering wants fewer "
        "restrictions on model access. Leadership wants both."
    )
    add_body(doc,
        "I'm an enterprise software engineer who also runs personal AI projects on my own Claude "
        "and Gemini subscriptions — a knowledge distillation pipeline, a KB quiz engine, and a "
        "process-monitoring MCP server. That combination gives me a useful vantage point: I feel "
        "the cost personally, I see the productivity question professionally, and I've had to find "
        "real answers for both. Here's what I've learned."
    )

    # Core mental model
    set_heading(doc, "The Core Mental Model: Route by Difficulty", 1)
    add_body(doc,
        "Stop treating your AI stack as a single model and start treating it as a fleet with a "
        "routing layer. Premium models (Claude Opus, Gemini Ultra) are best for complex multi-step "
        "reasoning, novel synthesis across sources, and high-stakes output where quality directly "
        "affects users. Everything else is a candidate for a cheaper or local alternative."
    )

    # 1. Multi-agent
    set_heading(doc, "1. Multi-Agent Pipelines — Match the Model to the Step", 1)
    add_body(doc,
        "Most workflows aren't a single task — they're a sequence of steps with very different "
        "complexity profiles. Assigning one premium model to the entire pipeline is like hiring a "
        "specialist surgeon to also do your hospital's paperwork. The better approach: decompose "
        "your workflow into steps, then assign the cheapest model capable of each step."
    )

    add_image(doc, diagram_paths["agent_pipeline"], caption="Figure 1 — Multi-Agent Pipeline: model assignment per step")

    add_body(doc, "Example — a document Q&A skill:")
    add_code_block(doc,
        "Step 1: Parse & clean raw input          → local model (Phi-3, Mistral 7B)\n"
        "Step 2: Classify intent / route query    → small cloud model (Haiku, Flash)\n"
        "Step 3: Retrieve relevant chunks (RAG)   → embedding model only, no LLM\n"
        "Step 4: Synthesize answer from chunks    → premium model (Claude Sonnet/Opus)\n"
        "Step 5: Format & validate output schema  → small cloud model (Haiku, Flash)"
    )

    add_body(doc, "Practical implementation tips:")
    add_bullet(doc, " Define a schema contract between agents — the output of each step is the structured input to the next.", "Schema contracts:")
    add_bullet(doc, " Each agent should have a single responsibility. Agents that do too much become expensive and opaque.", "Single responsibility:")
    add_bullet(doc, " Build confidence scoring into classification steps — if the cheaper model isn't confident, escalate.", "Confidence scoring:")
    add_bullet(doc, " With MCP support, wire different model-backed agents as distinct MCP servers, each domain-specialized.", "MCP integration:")
    doc.add_paragraph()

    # 2. Prompt caching
    set_heading(doc, "2. Prompt/Token Caching — The Fastest Win", 1)
    add_body(doc,
        "Claude's prompt caching lets you cache a prefix of your context (system prompt, documents, "
        "examples) and pay approximately 10% of the input token cost on repeated calls that share "
        "that prefix."
    )

    add_image(doc, diagram_paths["prompt_caching"], caption="Figure 2 — Prompt Caching: cost comparison per query")

    add_body(doc, "Rules for getting cache hits:")
    add_bullet(doc, "Put stable content first (system prompt, docs, few-shot examples).")
    add_bullet(doc, "Put variable content last (the user's actual question).")
    add_bullet(doc, "Use consistent cache_control breakpoints across requests in the same session.")
    add_bullet(doc, "Cache TTL is 5 minutes on Anthropic's API — keep call cadence inside that window or re-warm explicitly.")
    doc.add_paragraph()
    add_body(doc,
        "In one of my personal projects — a knowledge base quiz engine that repeatedly queries the "
        "same document corpus — this single change reduced input token costs by over 80%."
    )

    # 3. Vector DBs
    set_heading(doc, "3. Vector DBs — Ground the Model in What It Cannot Know", 1)
    add_body(doc,
        "Before getting to the cost angle, it's worth understanding why a vector database is "
        "necessary in the first place. Large language models reason from their latent space — "
        "knowledge compressed into weights during training. That works well for general concepts, "
        "but it fails in predictable ways:"
    )
    add_bullet(doc, " The model was trained on docs from months ago. It will confidently describe a method signature that no longer exists.", "Version-specific SDK and API references:")
    add_bullet(doc, " Windows kernel internals, IRQL rules, minifilter callback contracts. Models have surface-level familiarity, not the precision required when the details actually matter.", "Deep domain knowledge:")
    add_bullet(doc, " Your component design, your team's architectural decisions, your internal APIs. The model has zero knowledge of these by definition.", "Internal project context:")

    add_body(doc,
        "The fix is RAG — injecting authoritative, current context into the prompt so the model "
        "reasons from your ground truth rather than its best guess."
    )

    add_image(doc, diagram_paths["rag_flow"], caption="Figure 3 — RAG Flow: retrieval-augmented generation pipeline")

    add_body(doc, "Sizing guidance:")
    add_bullet(doc, "Top-K = 3–5 chunks is usually enough. More is rarely better and always more expensive.")
    add_bullet(doc, "Chunk size: 512–1024 tokens per chunk gives retrieval precision without truncating context.")
    add_bullet(doc, "Re-rank before sending: use a lightweight cross-encoder or BM25 hybrid.")
    doc.add_paragraph()
    add_body(doc,
        "In my personal process-monitoring project, instead of sending full logs to Claude for every "
        "alert, I embed process events, retrieve only the anomalous window, and send that. Latency "
        "dropped, cost dropped — and critically, the model stopped hallucinating event fields that "
        "didn't exist in the actual log schema."
    )
    add_body(doc,
        "Key reframe: RAG isn't primarily a cost optimization — it's a correctness optimization "
        "that also happens to reduce cost. That's a much easier sell in an enterprise conversation "
        "where accuracy matters more than invoices."
    )

    # 4. Local models
    set_heading(doc, "4. Local Models — Free the Routine Work", 1)
    add_body(doc,
        "Not every task needs a frontier model. Local models (Ollama + Llama 3, Mistral, Phi-3) "
        "running on your own hardware handle a surprising amount of routine work."
    )

    add_table(doc,
        ["Task", "Use Local", "Use Premium"],
        [
            ["Intent classification", "✓", ""],
            ["Entity extraction (structured schema)", "✓", ""],
            ["Routing / triage", "✓", ""],
            ["Summarization of well-structured input", "✓", ""],
            ["Complex reasoning across multiple documents", "", "✓"],
            ["Code generation with non-trivial logic", "", "✓"],
            ["Novel synthesis, edge cases, ambiguity", "", "✓"],
        ]
    )

    add_body(doc,
        "My personal knowledge distillation pipeline uses a local 7B model to pre-process and "
        "structure raw input, then sends only the structured output to Claude for synthesis. Token "
        "input to Claude drops 60–70% because the local model did the cleanup first."
    )

    # 5. Context window discipline
    set_heading(doc, "5. Context Window Discipline", 1)
    add_body(doc,
        "Premium models have large context windows. That doesn't mean you should fill them. "
        "Cost scales linearly with input tokens. Quality does not."
    )
    add_bullet(doc, " Summarize prior conversation turns rather than appending them verbatim.", "Compress before sending:")
    add_bullet(doc, " Ask the model to return JSON with a defined schema — smaller, parseable outputs, no rambling preambles.", "Structured output contracts:")
    add_bullet(doc, " Three high-quality examples beat ten mediocre ones — and cost 70% less.", "Trim few-shot examples:")
    add_bullet(doc, " Remove headers, footers, page numbers, and formatting artifacts from documents before sending.", "Strip metadata:")

    # 6. Programmable Tool Calling
    set_heading(doc, "6. Programmable Tool Calling — The Model as Orchestrator", 1)
    add_body(doc,
        "Tool calling (function calling) is commonly seen as a way to connect models to external "
        "data. Its bigger value is architectural: it lets the premium model act as an orchestrator "
        "that decides what to invoke — and what not to invoke — rather than doing all reasoning inline."
    )
    add_body(doc, "Three efficiency patterns emerge from this:")
    add_bullet(doc,
        " The model can call vector_search() and classify_intent() in a single response. "
        "Instead of two sequential round trips — each paying full input token cost — you get one. "
        "For pipelines with multiple retrieval steps this compounds quickly.",
        "Parallel tool calls reduce round trips:"
    )
    add_bullet(doc,
        " Define tools that wrap your local model, embedding service, or cache lookup. "
        "The premium model decides whether to invoke them based on the task. A high-confidence "
        "classification routes to run_local_model(); a complex synthesis stays in-model. "
        "The model becomes the router — no separate routing layer needed.",
        "Tools as cheap-service wrappers:"
    )
    add_bullet(doc,
        " Instead of prompt-engineering 'return valid JSON with fields X, Y, Z' and then "
        "parsing brittle free text, define a tool with a strict JSON schema. The model is "
        "constrained to call it correctly. Output token count drops, parsing reliability goes to "
        "near-100%, and you eliminate a whole class of retry logic.",
        "Schema enforcement via tool definitions:"
    )

    add_image(doc, diagram_paths["tool_calling"], caption="Figure 5 — Programmable Tool Calling: model as orchestrator")

    add_body(doc, "Practical tips:")
    add_bullet(doc, "Keep tool descriptions concise — they count against your input tokens on every call.")
    add_bullet(doc, "Group tools logically: retrieval tools, transformation tools, output tools. Fewer tools in scope = less decision overhead for the model.")
    add_bullet(doc, "Use tool_choice: 'required' when you always need structured output — prevents the model from answering in prose instead.")
    add_bullet(doc, "Log every tool invocation with latency and token counts. This is where you find the expensive surprises.")
    doc.add_paragraph()

    # 7. Async batching
    set_heading(doc, "7. Async Batching — Don't Pay Rush Rates for Batch Work", 1)
    add_body(doc,
        "Anthropic's Message Batches API gives you 50% off input and output tokens for non-real-time "
        "workloads with 24-hour turnaround. Nightly evaluations, bulk document processing, quiz "
        "generation pipelines — none of these should hit the synchronous API. Same principle applies "
        "across providers: identify workloads that can tolerate latency and route them to batch endpoints."
    )

    # Full stack
    set_heading(doc, "The Full Stack", 1)
    add_body(doc,
        "Each layer below reduces what reaches the premium model. The premium model does what only "
        "it can do — and nothing else."
    )
    add_image(doc, diagram_paths["full_stack"], caption="Figure 6 — Full Efficient AI Stack: combined pipeline")

    # Enterprise debate
    set_heading(doc, "Answering the Enterprise Debate", 1)
    add_body(doc,
        "When the cost-vs-productivity question comes up at work, the honest answer is: the question "
        "is usually posed too broadly to be useful. It's not 'is Claude worth it?' — it's 'which "
        "steps in our workflows justify the cost, and which are we routing there out of habit?'"
    )
    add_body(doc, "Track these metrics per feature:")
    add_bullet(doc, "Input tokens per task — primary cost driver")
    add_bullet(doc, "Cache hit rate — target >70% for KB-heavy workloads")
    add_bullet(doc, "Local/mid-tier deflection rate — % of steps resolved without a premium call")
    add_bullet(doc, "Premium model invocation rate — how often does the expensive agent actually fire?")
    add_bullet(doc, "Output token ratio — output/input; high ratios mean you're extracting value")

    # Mindset shift
    set_heading(doc, "The Mindset Shift", 1)
    add_body(doc,
        "Premium models are exceptional tools. Routing every step of every workflow through them is "
        "like staffing a team entirely with principal engineers when juniors, mids, and seniors each "
        "have work suited to their level."
    )
    add_body(doc,
        "The engineers extracting real value from frontier AI are the ones who build agent pipelines "
        "with deliberate model assignment — so that when the premium model runs, it runs on exactly "
        "the right input, at exactly the right step, at exactly the right time."
    )
    add_body(doc, "Build the infrastructure. Make every premium token count.")

    doc.add_paragraph()
    closing = doc.add_paragraph(
        "I'm an enterprise software engineer who builds AI tooling and Windows systems as personal "
        "projects. Thoughts on the cost-vs-productivity debate or multi-agent patterns? I'd like to "
        "hear how others are approaching this."
    )
    closing.runs[0].italic = True

    out_path = OUT_DIR / "efficient_ai_usage.docx"
    doc.save(str(out_path))
    print(f"\nDocument saved: {out_path}")
    return out_path


if __name__ == "__main__":
    diagram_paths = render_diagrams()
    build_doc(diagram_paths)
