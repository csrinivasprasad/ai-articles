import subprocess
import os
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = Path(__file__).parent
PLANTUML_JAR = r"C:\tools\plantuml.jar"

# ---------------------------------------------------------------------------
# PlantUML sources
# ---------------------------------------------------------------------------

DIAGRAMS = {
    "graph_basic": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title A Graph: Nodes and Edges

(A) -- (B)
(A) -- (C)
(B) -- (D)
(C) -- (D)
@enduml
""",

    "graph_directed": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title Directed Graph

(A) --> (B)
(B) --> (C)
(C) --> (D)
@enduml
""",

    "graph_dag": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title DAG — Directed Acyclic Graph (Pipeline)

(A) --> (B)
(A) --> (C)
(B) --> (D)
(C) --> (D)
@enduml
""",

    "graph_cyclic": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title Cyclic Directed Graph (Agent Loop)

(Agent) --> (Tool)
(Tool) --> (Agent) : observation
(Agent) --> (End) : no more tool calls
@enduml
""",

    "repl": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title REPL Pattern

:User Input;
repeat
  :LLM Call;
  :Output to User;
repeat while (more input?) is (yes)
->no;
:End;
@enduml
""",

    "react": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title ReAct — Reasoning + Acting

start
:User Query;
repeat
  :Thought\\n(LLM reasons about next step);
  :Action\\n(select tool + arguments);
  :Observation\\n(tool result returned to LLM);
repeat while (final answer reached?) is (no)
->yes;
:Output Final Answer;
stop
@enduml
""",

    "plan_execute": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title Plan-and-Execute Pattern

start
:Planner LLM\\n→ generates [Step 1, Step 2, Step 3];
:Executor receives plan;
repeat
  :Execute current step;
  :Update state with result;
  if (Goal achieved?) then (yes)
    break
  else (no)
    if (Step failed?) then (yes)
      :Re-planner LLM\\n→ revise remaining plan;
    endif
  endif
repeat while (steps remain)
:Output final result;
stop
@enduml
""",

    "reflection": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title Reflection / Self-Critique Pattern

start
:User Task;
repeat
  :Generator LLM\\n→ produces draft output;
  :Critic LLM\\n→ evaluates against criteria;
  :Critique added to context;
repeat while (quality gate passed?) is (no)
->yes;
:Return final output;
stop
@enduml
""",

    "multi_agent": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title Multi-Agent — Orchestrator + Subagents

start
:Orchestrator LLM\\n→ decomposes task;

fork
  :Research Agent\\n(web search, RAG);
fork again
  :Code Agent\\n(write + execute code);
fork again
  :Review Agent\\n(quality check);
end fork

:Orchestrator\\n→ synthesize results;
:Final Output;
stop
@enduml
""",

    "tot": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title Tree of Thoughts

start
:Root Thought;
fork
  :Branch A;
  if (Promising?) then (no)
    :Prune;
    kill
  else (yes)
    :Branch A1 → Solution;
  endif
fork again
  :Branch B;
  if (Promising?) then (yes)
    fork
      :Branch B1 → dead end;
      kill
    fork again
      :Branch B2 → Solution;
    end fork
  else (no)
    :Prune;
    kill
  endif
end fork
:Select best solution;
stop
@enduml
""",

    "memory_agent": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title Memory-Augmented Agent

rectangle "Memory Stores" {
  database "In-Context\\n(chat history)" as CTX
  database "Episodic\\n(vector DB / RAG)" as EPI
  database "Semantic\\n(entity KB)" as SEM
  database "Procedural\\n(tool-use patterns)" as PRO
}

rectangle "Agent Loop" {
  component "Query\\nMemory" as QM
  component "LLM\\nReasoning" as LLM
  component "Act /\\nUpdate Memory" as ACT
}

QM --> LLM
LLM --> ACT
ACT --> QM

CTX <--> QM
EPI <--> QM
SEM <--> QM
PRO <--> ACT
@enduml
""",

    "event_driven": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title Event-Driven Agentic Loop

queue "Event Sources\\n(webhook, cron,\\nfile change, ETW)" as EVT
component "Agent Runtime" as AGT
database "Persistent State" as ST
queue "Downstream\\nSystems" as DS

EVT --> AGT : trigger
AGT <--> ST : read/write state
AGT --> DS : structured output
@enduml
""",

    "langgraph_core": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title LangGraph Core — Stateful Graph

[*] --> AgentNode : invoke(input, config)
AgentNode --> ToolNode : tool_call in output
ToolNode --> AgentNode : observation (loop)
AgentNode --> [*] : no tool_call → END

note right of AgentNode
  State dict flows through
  every node. Checkpointed
  after each execution.
end note
@enduml
""",

    "langgraph_hitl": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title LangGraph — Human-in-the-Loop

start
:Agent runs to interrupt_before node;
:State checkpointed to Postgres;
:PAUSE — surface state to human;
:Human reviews + optionally edits state;
:graph.invoke(None, config) — resume;
:Execution continues from checkpoint;
stop
@enduml
""",

    "langgraph_mapreduce": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title LangGraph — Map-Reduce via Send API

start
:Split Node\\n→ generates N work items;

fork
  :Send(task_1)\\nworker branch;
fork again
  :Send(task_2)\\nworker branch;
fork again
  :Send(task_N)\\nworker branch;
end fork

:Aggregate Node\\n→ merge all results;
:Final Output;
stop
@enduml
""",

    "framework_spectrum": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title Determinism vs Emergence Spectrum

left to right direction

rectangle "DETERMINISTIC" as D #D5E8D4
rectangle "LangGraph\\n(you define flow)" as LG #D5E8D4
rectangle "CrewAI\\n(task pipeline)" as CA #FFE6CC
rectangle "AutoGen\\n(emergent conversation)" as AU #F8CECC
rectangle "EMERGENT" as E #F8CECC

D --> LG
LG --> CA
CA --> AU
AU --> E
@enduml
""",

    "checkpoint_arch": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title LangGraph — Production Checkpoint Architecture

actor "API Request" as REQ
component "LangGraph App\\n(stateless)" as APP #D5E8D4
database "PostgreSQL\\nCheckpoint Store" as PG #E8F4FD

REQ --> APP : invoke(input, config{thread_id})
APP --> PG : write checkpoint after each node
APP <-- PG : read checkpoint on resume
REQ <-- APP : final output

note right of PG
  thread_id, checkpoint_id,
  parent_id, timestamp,
  channel_values (full state),
  metadata (step, writes)
end note
@enduml
""",

    "resume_flow": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title LangGraph — Resume After Failure

start
:invoke(input, thread_id="job-99");
:Node A executes → checkpoint saved;
:Node B executes → checkpoint saved;
:Node C crashes — EXCEPTION;
note right : Last checkpoint = end of Node B

:... process restarts ...;

:invoke(None, thread_id="job-99");
note right : input=None signals resume
:Load checkpoint — skip A and B;
:Re-execute Node C from saved state;
:Continue D, E → complete;
stop
@enduml
""",

    "time_travel": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title LangGraph — Time-Travel / Branching

object "CP-1" as c1
object "CP-2" as c2
object "CP-3" as c3
object "CP-4" as c4
object "CP-5" as c5

object "CP-3 (fork)" as c3f
object "CP-4'" as c4x
object "CP-5'" as c5x

c1 --> c2 : original run
c2 --> c3
c3 --> c4
c4 --> c5

c3 ..> c3f : update_state()\\ncreates fork
c3f --> c4x : diverges here
c4x --> c5x
@enduml
""",

    "state_machine": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title State Machine — States and Transitions

[*] --> Idle : start
Idle --> Processing : input received
Processing --> AwaitingTool : tool call needed
AwaitingTool --> Processing : tool result returned
Processing --> Done : final answer ready
Done --> [*]

note right of Processing
  Only ONE state is
  active at a time.
  The state dict holds
  all accumulated data.
end note
@enduml
""",

    "mapreduce": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title MapReduce — Fan-Out and Fan-In

start
:Input dataset or task list;
:MAP phase — split into N chunks\n(fan-out: work distributed in parallel);

fork
  :Worker 1\\nprocesses chunk 1;
fork again
  :Worker 2\\nprocesses chunk 2;
fork again
  :Worker N\\nprocesses chunk N;
end fork

:REDUCE phase — aggregate results\\n(fan-in: partial results merged);
:Final output;
stop
@enduml
""",

    "vector_embeddings": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title Vector Embeddings — Semantic Search Pipeline

start
:"How do I use LangGraph?"\\n(query text);
:Embedding model\\n→ converts text to vector\\n[0.12, -0.83, 0.45, ...\\n1536 numbers];
:Vector DB similarity search\\n(cosine distance to all stored vectors);
note right
  Similar meaning
  → similar direction
  in high-dimensional space
end note
:Top-K most similar chunks\\n(only relevant context);
:LLM receives focused context\\n→ grounded answer;
stop
@enduml
""",

    "project_map": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title Project Pattern Map

package "Raw Claude API" #D5E8D4 {
  component "ai-kb-quiz\\n(RAG + Hybrid Routing + REPL)" as KBQ
  component "ai-kd\\n(ReAct Tool Calling)" as KD
}

package "LangGraph + Pydantic AI" #FFE6CC {
  component "ai-threat-state-graph\\n(Stateful Detection Graph)" as TSG
}

package "MCP Server" #E8F4FD {
  component "ai-procwatch-mcp\\n(Tool Registry + Two-Tier Triage)" as PWM
}

package "Deterministic — No LLM" #F0F0F0 {
  component "ai-asset-sweeper\\n(Lua DSL Scripting)" as ASW
}
@enduml
""",

    "kb_quiz_flow": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title ai-kb-quiz — Hybrid Routing + PTC + RAG

start
:Markdown KB files;
:Chunker (300-600 chars);
:SHA-256 dedup → Context Cache;
:Embedder (Ollama or sentence-transformers);
:ChromaDB Vector Store;

note right : Indexing path (offline)

:User query;
:Retriever → top-k chunks;
:PTC Compression\\n(compress before model call);
:Router\\n(complexity threshold);
fork
  :LocalAdapter\\n(Ollama);
fork again
  :PremiumAdapter\\n(Claude Sonnet);
end fork
:Scorer + follow-up suggestions;
:Interactive REPL loop;
stop
@enduml
""",

    "ai_kd_flow": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title ai-kd — ReAct Tool Calling Loop

start
:Debug session starts\\n(CdbSession CLI or pykd);
:Initial context sent to Claude Sonnet;

repeat
  :Claude: Thought + Action\\n(WinDbg command);
  :Sanitiser strips kernel addresses;
  :execute_windbg_action\\n(run in CdbSession);
  :Sanitiser strips addresses from output;
  :Observation returned to Claude;
repeat while (root cause found\\nOR max 25 iters?) is (no)
->yes;

:AnalysisResult\\n(root-cause + tool_call_log);
stop
@enduml
""",

    "threat_state_flow": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title ai-threat-state-graph — CAAD Loop

start
:edr-sensor.sys\\n(Ring 0: callbacks → MPSC ring buffer);
:edr-service.exe\\n(IOCP Proactor, IRP bridge);
:ThreatContext created;

repeat
  :COLLECT → gather telemetry;
  :ANALYZE\\nYARA-X | Sigma | Lua | PS | Python;
  :ACT → verdict candidate;
  :DECIDE → confidence check;
repeat while (confidence < threshold\\nAND iters < 8) is (yes)
->no;

if (Phase 2 LLM escalation?) then (yes)
  :Claude Sonnet — advisory only,\\noff critical path;
endif

:Verdict (classification, MITRE, IOCs);
stop
@enduml
""",

    "procwatch_flow": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title ai-procwatch-mcp — Two-Tier Triage

start

fork
  :KrabsETW\\n(Process, File, Network,\\nRegistry, VirtualAlloc);
fork again
  :Windows eBPF\\n(socket, DKOM detection);
fork again
  :NTFS USN Journal\\n(file create/rename/delete);
end fork

:OCSF 1.5 + STIX 2.1 Normalization;
:Behavioral Genome (sqlite-vec);

:MCP Tool Registry;

:Tier 1: Ollama triage\\n(qwen3:4b — fast, local);

if (High-confidence anomaly?) then (yes)
  :Tier 2: Claude escalation\\n(chunked genome submission);
endif

:Verdict JSON\\n(verdict, score, MITRE, reasons);
stop
@enduml
""",

    "asset_sweeper_flow": """
@startuml
!theme plain
skinparam backgroundColor #FAFAFA
skinparam ArrowColor #444444
title ai-asset-sweeper — Sandboxed Lua DSL

start
:Orchestrator discovers *.lua scripts;

repeat
  :LuaEngine — sandboxed lua_State\\n(os.execute removed etc.);

  fork
    :probe.registry;
  fork again
    :probe.filesystem;
  fork again
    :probe.credentials;
  fork again
    :probe.modelheader;
  fork again
    :probe.process;
  end fork

  :Accumulate findings\\n(asset_type, severity, location);
repeat while (more scripts?) is (yes)
->no;

:Reporter (Text / JSON / CSV);
stop
@enduml
""",
}


# ---------------------------------------------------------------------------
# Word document helpers  (same as generate_article.py)
# ---------------------------------------------------------------------------

def set_heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


def add_code_block(doc, code):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x1E)
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), "F0F0F0")
    p._p.get_or_add_pPr().append(shading)
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_image(doc, path, width_inches=5.5, caption=None):
    if not Path(path).exists():
        doc.add_paragraph(f"[Diagram not rendered: {path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    if caption:
        cp = doc.add_paragraph(caption)
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.runs[0].italic = True
        cp.runs[0].font.size = Pt(9)


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        hdr_cells[i].paragraphs[0].runs[0].bold = True
        shading = OxmlElement("w:shd")
        shading.set(qn("w:val"), "clear")
        shading.set(qn("w:color"), "auto")
        shading.set(qn("w:fill"), "D5E8D4")
        hdr_cells[i]._tc.get_or_add_tcPr().append(shading)
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = val
    doc.add_paragraph()


# ---------------------------------------------------------------------------
# Render diagrams
# ---------------------------------------------------------------------------

def render_diagrams():
    print("Rendering PlantUML diagrams...")
    paths = {}
    for name, src in DIAGRAMS.items():
        puml_path = OUT_DIR / f"agpat_{name}.puml"
        png_path  = OUT_DIR / f"agpat_{name}.png"
        puml_path.write_text(src.strip(), encoding="utf-8")
        result = subprocess.run(
            ["java", "-jar", PLANTUML_JAR, str(puml_path), "-o", str(OUT_DIR)],
            capture_output=True, text=True
        )
        if result.returncode != 0:
            print(f"  ERROR {name}: {result.stderr}")
        else:
            print(f"  OK: agpat_{name}.png")
        paths[name] = png_path
    return paths


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

def build_doc(dp):  # dp = diagram_paths dict
    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("AI Agent Patterns", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("Fundamentals to Advanced — with Framework Comparisons and Real Project Analysis")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].italic = True
    sub.runs[0].font.size = Pt(13)
    doc.add_paragraph()

    add_body(doc,
        "This document covers every major AI agent execution pattern — from the simplest REPL loop "
        "through production-grade stateful graph systems — alongside a detailed comparison of "
        "LangGraph, AutoGen, and CrewAI, and an analysis of how these patterns appear in real projects."
    )

    # ── Part 0: Graph Fundamentals ───────────────────────────────────────────
    set_heading(doc, "Part 0 — Graph Fundamentals", 1)
    add_body(doc,
        "A graph is a data structure made of two things: nodes (the things) and edges (the "
        "connections between things). Understanding graph types is essential before exploring "
        "agent frameworks, because most modern agent runtimes are graph execution engines."
    )

    add_image(doc, dp["graph_basic"], width_inches=3.5, caption="Figure 0.1 — Basic undirected graph")

    set_heading(doc, "Types of Graphs", 2)

    add_body(doc, "Directed graph — edges have a direction. 'A calls B' is not the same as 'B calls A'.")
    add_image(doc, dp["graph_directed"], width_inches=3.5, caption="Figure 0.2 — Directed graph")

    add_body(doc,
        "DAG (Directed Acyclic Graph) — directed, no cycles. Used for pipelines where execution "
        "never goes back to a previous step."
    )
    add_image(doc, dp["graph_dag"], width_inches=3.5, caption="Figure 0.3 — DAG: parallel paths, no cycles")

    add_body(doc,
        "Cyclic directed graph — has loops. Used in agent frameworks for the ReAct tool loop: "
        "the agent node and tool node loop back to each other until a condition is met."
    )
    add_image(doc, dp["graph_cyclic"], width_inches=4.0, caption="Figure 0.4 — Cyclic graph: agent loop")

    set_heading(doc, "Why Graphs for Agents?", 2)
    add_body(doc,
        "An agent workflow maps naturally onto a graph. Before graph-based frameworks, agent "
        "control flow was hidden inside LLM prompts or tangled if/else chains. Representing it "
        "as an explicit graph makes execution visible, testable, and modifiable without touching "
        "the LLM logic. The graph is the map of execution."
    )
    add_table(doc,
        ["Graph Concept", "Agent Equivalent"],
        [
            ["Node", "Unit of work (LLM call, tool call, router, human approval)"],
            ["Directed edge", "After this, run that"],
            ["Conditional edge", "Go left if X, go right if Y"],
            ["Cycle", "Keep looping until condition is met"],
            ["DAG", "Linear pipeline, no re-execution"],
        ]
    )

    # State Machines
    set_heading(doc, "State Machines", 2)
    add_body(doc,
        "A state machine is a system that can be in exactly one state at a time, with defined "
        "rules for transitioning between states. It has three core concepts:"
    )
    add_bullet(doc, "a finite set of conditions the system can be in.", "States —")
    add_bullet(doc, "events or conditions that move the system from one state to another.", "Transitions —")
    add_bullet(doc, "the state the system is in right now.", "Current state —")
    doc.add_paragraph()
    add_image(doc, dp["state_machine"], width_inches=5.0, caption="Figure 0.5 — State machine: one active state at a time")
    add_body(doc,
        "LangGraph is a state machine runtime. Its nodes are states, its edges are transitions, "
        "and its typed state dict is the current state — carrying all accumulated data as execution "
        "moves through the graph. The checkpointer snapshots the current state after every transition "
        "so it can be resumed or rewound."
    )

    # MapReduce / Fan-out + Fan-in
    set_heading(doc, "MapReduce — Fan-Out and Fan-In", 2)
    add_body(doc,
        "MapReduce is a programming model for processing large datasets in parallel, introduced by "
        "Google in 2004. It has two phases:"
    )
    add_bullet(doc, "split the input into independent chunks and process each in parallel (fan-out).", "Map —")
    add_bullet(doc, "collect all partial results and merge them into a single output (fan-in).", "Reduce —")
    doc.add_paragraph()
    add_image(doc, dp["mapreduce"], width_inches=5.0, caption="Figure 0.6 — MapReduce: parallel fan-out then fan-in aggregate")
    add_body(doc,
        "In agent frameworks this pattern appears whenever you have a list of independent work items "
        "that can be processed concurrently — document chunks, search queries, code files. "
        "LangGraph implements it via the Send API: the split node dispatches N parallel branches "
        "(fan-out), each runs independently, and an aggregate node merges the results (fan-in). "
        "Wall-clock time is bounded by the slowest worker, not the sum of all workers."
    )

    # Vector Embeddings
    set_heading(doc, "Vector Embeddings and Vector Databases", 2)
    add_body(doc,
        "An embedding is a list of numbers (a vector) that represents the meaning of a piece of "
        "text. An embedding model reads text and outputs a vector — typically 768 to 1536 numbers. "
        "The key property: text with similar meaning produces vectors that point in similar "
        "directions in that high-dimensional space. This is measured with cosine similarity."
    )
    add_body(doc,
        "A vector database stores these vectors and answers the question: "
        "'Given this query vector, which stored vectors are most similar?' "
        "This is semantic search — it finds relevant content even when the words don't match exactly."
    )
    add_image(doc, dp["vector_embeddings"], width_inches=5.0, caption="Figure 0.7 — Vector embeddings: text → vector → similarity search")
    add_body(doc,
        "Why this matters for agents: keyword search fails when the query uses different words than "
        "the document. Vector search finds conceptually related content regardless of exact wording. "
        "This is the foundation of RAG (Retrieval-Augmented Generation) and every memory-augmented "
        "agent pattern — inject only the relevant chunks into the prompt rather than the entire "
        "knowledge base, keeping context windows focused and token costs low."
    )
    add_table(doc,
        ["Term", "What it means"],
        [
            ["Embedding", "A vector (list of numbers) representing the meaning of text"],
            ["Cosine similarity", "How similar two vectors are — 1.0 = identical direction, 0 = unrelated"],
            ["Vector DB", "Database optimised for nearest-neighbour search over embeddings (ChromaDB, Pinecone, pgvector)"],
            ["Top-K", "The K most similar chunks returned by the search — typically 3-5"],
            ["RAG", "Retrieval-Augmented Generation — fetch relevant chunks, inject into prompt, LLM reasons from them"],
        ]
    )

    # ── Part 1: Foundational Patterns ────────────────────────────────────────
    set_heading(doc, "Part 1 — Foundational Agent Patterns", 1)

    # 1. REPL
    set_heading(doc, "1. REPL (Read-Eval-Print Loop)", 2)
    add_body(doc,
        "The simplest pattern. The agent reads input, calls the LLM, outputs a response, and loops. "
        "No memory, no tool use, no planning. Essentially a chatbot loop."
    )
    add_image(doc, dp["repl"], width_inches=3.5, caption="Figure 1.1 — REPL pattern")

    # 2. ReAct
    set_heading(doc, "2. ReAct (Reasoning + Acting)", 2)
    add_body(doc,
        "The dominant pattern for tool-using agents. The LLM alternates between Thought, Action, "
        "and Observation in a scratchpad loop until it reaches a final answer. Reasoning and acting "
        "are interleaved — each observation feeds back into the next thought, grounding the LLM "
        "in real results."
    )
    add_image(doc, dp["react"], width_inches=4.5, caption="Figure 1.2 — ReAct: interleaved reasoning and acting")

    # 3. Plan-and-Execute
    set_heading(doc, "3. Plan-and-Execute", 2)
    add_body(doc,
        "Separates planning from execution. A planner LLM generates a step-by-step plan upfront; "
        "an executor agent runs each step. Better for long-horizon tasks. A re-planner triggers "
        "if execution fails, which is the key differentiator from ReAct."
    )
    add_image(doc, dp["plan_execute"], width_inches=5.0, caption="Figure 1.3 — Plan-and-Execute with re-planning")

    # 4. Reflection
    set_heading(doc, "4. Reflection / Self-Critique", 2)
    add_body(doc,
        "The agent generates output, then a second LLM call critiques it, and the agent revises. "
        "Loops until a quality gate is passed."
    )
    add_image(doc, dp["reflection"], width_inches=4.5, caption="Figure 1.4 — Reflection loop")
    add_body(doc, "Variants:")
    add_bullet(doc, "stores failed attempts in memory to avoid repeating mistakes.", "Reflexion —")
    add_bullet(doc, "critique against a fixed set of principles.", "Constitutional AI —")
    doc.add_paragraph()

    # 5. Multi-Agent
    set_heading(doc, "5. Multi-Agent (Orchestrator + Subagents)", 2)
    add_body(doc,
        "One orchestrator decomposes a task and delegates to specialist subagents. Subagents can "
        "run in parallel. The orchestrator maintains shared state and synthesizes results."
    )
    add_image(doc, dp["multi_agent"], width_inches=4.5, caption="Figure 1.5 — Multi-Agent with parallel subagents")

    # 6. Tree of Thoughts
    set_heading(doc, "6. Tree of Thoughts (ToT)", 2)
    add_body(doc,
        "Extends chain-of-thought by exploring a tree of reasoning paths rather than one linear "
        "chain. The model evaluates intermediate steps and prunes dead branches (BFS or DFS). "
        "Best for combinatorial search problems: puzzles, theorem proving, multi-step math."
    )
    add_image(doc, dp["tot"], width_inches=5.0, caption="Figure 1.6 — Tree of Thoughts with pruning")

    # 7. Memory-Augmented
    set_heading(doc, "7. Memory-Augmented Agent", 2)
    add_body(doc,
        "Any pattern above extended with explicit memory stores. The agent reads from and writes "
        "to memory on each cycle, enabling long-term recall across sessions."
    )
    add_image(doc, dp["memory_agent"], width_inches=5.5, caption="Figure 1.7 — Memory-Augmented Agent: four memory stores")

    # 8. Event-Driven
    set_heading(doc, "8. Event-Driven Agentic Loop", 2)
    add_body(doc,
        "The agent is triggered by external events (webhooks, file changes, cron, ETW telemetry) "
        "rather than user input. Runs autonomously, emits structured outputs to downstream systems."
    )
    add_image(doc, dp["event_driven"], width_inches=5.0, caption="Figure 1.8 — Event-Driven loop")

    # Pattern selection table
    set_heading(doc, "Pattern Selection Heuristic", 2)
    add_table(doc,
        ["Scenario", "Best Pattern"],
        [
            ["Single-turn Q&A", "REPL"],
            ["Tool use, web search", "ReAct"],
            ["Long multi-step task", "Plan-and-Execute"],
            ["Quality-sensitive generation", "Reflection"],
            ["Parallelizable subtasks", "Multi-Agent"],
            ["Search / optimization problems", "Tree of Thoughts"],
            ["Long-term recall across sessions", "Memory-Augmented"],
            ["Persistent autonomous work", "Event-Driven"],
        ]
    )

    # ── Part 2: LangGraph ────────────────────────────────────────────────────
    set_heading(doc, "Part 2 — LangGraph Patterns In Depth", 1)
    add_body(doc,
        "LangGraph is built on one foundational abstraction — stateful directed graphs — and layers "
        "several agent patterns on top of it. The key design principle: control flow lives in your "
        "code, not in the LLM's output. This makes agent behavior explicit, inspectable, and testable."
    )

    set_heading(doc, "Core Abstraction — State Machine", 2)
    add_image(doc, dp["langgraph_core"], width_inches=5.5, caption="Figure 2.1 — LangGraph core: stateful graph with checkpointing")
    add_body(doc,
        "Nodes are Python functions (LLM calls, tool calls, logic). Edges are transitions "
        "(fixed or conditional). State is a typed dict that flows through every node and accumulates "
        "changes. It is checkpointed after each node execution."
    )

    set_heading(doc, "Human-in-the-Loop via Interrupt + Checkpointer", 2)
    add_image(doc, dp["langgraph_hitl"], width_inches=5.0, caption="Figure 2.2 — LangGraph HITL: pause, edit, resume")
    add_code_block(doc,
        "graph = workflow.compile(\n"
        "    checkpointer=checkpointer,\n"
        "    interrupt_before=[\"approve_action\"]  # pause before this node\n"
        ")\n"
        "graph.update_state(config, {\"approved\": True})  # human edits state\n"
        "graph.invoke(None, config)                        # resume"
    )

    set_heading(doc, "Map-Reduce via Send API", 2)
    add_image(doc, dp["langgraph_mapreduce"], width_inches=5.0, caption="Figure 2.3 — LangGraph Map-Reduce: parallel fan-out / fan-in")

    set_heading(doc, "LangGraph Pattern-to-Primitive Mapping", 2)
    add_table(doc,
        ["LangGraph Primitive", "Agent Pattern Enabled"],
        [
            ["Cycles", "ReAct, Reflection"],
            ["Conditional edges", "Plan-and-Execute, Supervisor"],
            ["Send API (parallel)", "Map-Reduce, Multi-Agent fan-out"],
            ["Interrupt + Checkpointer", "Human-in-the-Loop, Resume"],
            ["Subgraphs", "Multi-Agent composition"],
        ]
    )

    # ── Part 3: Framework Comparison ─────────────────────────────────────────
    set_heading(doc, "Part 3 — Framework Comparison: LangGraph vs AutoGen vs CrewAI", 1)

    add_table(doc,
        ["Framework", "Core Abstraction", "Mental Model"],
        [
            ["LangGraph", "Stateful directed graph", "You are the architect — explicit nodes, edges, state"],
            ["AutoGen", "Conversational agents", "Agents are actors that talk to each other"],
            ["CrewAI", "Role-based crew", "Agents are employees with job descriptions"],
        ]
    )

    set_heading(doc, "Determinism vs Emergence Spectrum", 2)
    add_image(doc, dp["framework_spectrum"], width_inches=5.5, caption="Figure 3.1 — Determinism spectrum: LangGraph → CrewAI → AutoGen")

    set_heading(doc, "Multi-Agent Topology Support", 2)
    add_table(doc,
        ["Topology", "LangGraph", "AutoGen", "CrewAI"],
        [
            ["Sequential pipeline", "Graph edges", "initiate_chat chain", "Sequential process"],
            ["Supervisor → workers", "Supervisor subgraph", "GroupChat with selector", "Hierarchical process"],
            ["Peer-to-peer conversation", "Manual graph cycle", "Native (core feature)", "Not native"],
            ["Parallel fan-out", "Send API", "Async agents", "Partial (async tasks)"],
            ["Nested subgraphs", "Native", "Nested GroupChat", "Nested crews (v0.8+)"],
        ]
    )

    set_heading(doc, "Human-in-the-Loop", 2)
    add_table(doc,
        ["Capability", "LangGraph", "AutoGen", "CrewAI"],
        [
            ["Interrupt before node", "Native (interrupt_before)", "human_input_mode on agent", "Limited"],
            ["Approve tool call", "Native", "human_input_mode=ALWAYS", "Not native"],
            ["Resume from checkpoint", "Native", "Not built-in", "Not built-in"],
            ["Edit state mid-run", "Native", "Not built-in", "Not built-in"],
        ]
    )

    set_heading(doc, "When to Use Which Framework", 2)
    add_table(doc,
        ["Use Case", "Best Fit", "Why"],
        [
            ["Production workflow with approvals", "LangGraph", "Checkpointing, interrupt, audit trail"],
            ["Data science / coding assistant", "AutoGen", "Native code execution sandbox"],
            ["RAG pipeline, document processing", "LangGraph", "Map-reduce, parallel fan-out"],
            ["Quick prototype multi-agent system", "CrewAI", "Minimal boilerplate, role-based intuition"],
            ["Research / experimental agents", "AutoGen", "Emergent conversation, flexible"],
            ["Long-running background automation", "LangGraph", "Persistent state, resume on failure"],
            ["Regulated / auditable systems", "LangGraph", "Time-travel and state snapshots"],
        ]
    )

    # ── Part 4: LangGraph Production Persistence ──────────────────────────────
    set_heading(doc, "Part 4 — LangGraph State Persistence in Production", 1)
    add_body(doc,
        "LangGraph persists state via a checkpointer attached to the compiled graph. After every "
        "node execution, the full state dict is serialized and saved. The graph becomes resumable "
        "from any point. Every run is identified by a thread_id."
    )

    set_heading(doc, "Production Checkpoint Architecture", 2)
    add_image(doc, dp["checkpoint_arch"], width_inches=5.5, caption="Figure 4.1 — Checkpoint architecture: stateless app + Postgres store")

    set_heading(doc, "Checkpointer Backends", 2)
    add_table(doc,
        ["Backend", "Use Case", "Notes"],
        [
            ["MemorySaver", "Dev / testing only", "Lost on process restart"],
            ["SqliteSaver", "Single-process production", "File-based, simple"],
            ["PostgresSaver", "Multi-process production", "Recommended for scale"],
            ["RedisSaver", "High-throughput", "TTL support"],
            ["Custom", "Any store", "Implement BaseCheckpointSaver"],
        ]
    )

    set_heading(doc, "Resume After Failure", 2)
    add_image(doc, dp["resume_flow"], width_inches=5.0, caption="Figure 4.2 — Resume: nodes already completed are not re-run")

    set_heading(doc, "Time-Travel and Branching", 2)
    add_image(doc, dp["time_travel"], width_inches=5.5, caption="Figure 4.3 — Time-travel: fork from any past checkpoint")
    add_body(doc, "Use cases:")
    add_bullet(doc, "Debug production failures by replaying with modified input.")
    add_bullet(doc, "A/B test different agent decisions from a shared starting point.")
    add_bullet(doc, "Recover from bad agent decisions without restarting from scratch.")
    doc.add_paragraph()

    set_heading(doc, "Production Gotchas", 2)
    add_table(doc,
        ["Gotcha", "Detail"],
        [
            ["State size bloat", "Store S3 keys / DB IDs in state, not raw document content"],
            ["Serialization", "Everything in state must be JSON-serializable; Pydantic models work"],
            ["Thread ID design", "Use domain IDs (order-123) not random UUIDs for external lookups"],
            ["Checkpoint retention", "Checkpoints accumulate forever — add TTL policy or cleanup job"],
            ["Idempotency", "Re-run nodes must be idempotent; LLM calls are not naturally idempotent"],
        ]
    )

    # ── Part 5: Real Projects ─────────────────────────────────────────────────
    set_heading(doc, "Part 5 — Patterns in Practice: Real Projects", 1)
    add_body(doc,
        "The following projects demonstrate how these patterns appear in real production and "
        "personal systems — from raw Claude API tool loops through LangGraph stateful graphs "
        "to fully deterministic Lua-scripted engines."
    )

    add_image(doc, dp["project_map"], width_inches=5.5, caption="Figure 5.0 — Project pattern map")

    # ai-kb-quiz
    set_heading(doc, "ai-kb-quiz — Hybrid Routing + PTC + RAG", 2)
    add_body(doc,
        "Patterns: Hybrid model routing, Process-Then-Communicate (PTC), Vectorized RAG, "
        "Interactive REPL learning loop."
    )
    add_image(doc, dp["kb_quiz_flow"], width_inches=5.5, caption="Figure 5.1 — ai-kb-quiz architecture")
    add_body(doc,
        "Key insight: The PTC pattern (Process-Then-Communicate) ensures raw KB content never "
        "hits the model — it is compressed first. The router avoids premium API calls for "
        "simple questions, routing them to a local Ollama model instead."
    )

    # ai-kd
    set_heading(doc, "ai-kd — ReAct Tool Calling (Kernel Debugger)", 2)
    add_body(doc,
        "Pattern: ReAct with a single tool (execute_windbg_action), plus a sanitization layer "
        "that strips kernel addresses before sending output to Claude."
    )
    add_image(doc, dp["ai_kd_flow"], width_inches=5.5, caption="Figure 5.2 — ai-kd: ReAct loop with sanitization")
    add_body(doc,
        "Key insight: The sanitization layer is a security boundary — it prevents prompt injection "
        "via malicious kernel output and avoids leaking KASLR addresses to the external API."
    )

    # ai-threat-state-graph
    set_heading(doc, "ai-threat-state-graph — CAAD Loop + Five Detection Engines", 2)
    add_body(doc,
        "Pattern: Stateful LangGraph graph with CAAD loop (Collect → Analyze → Act → Decide). "
        "Five deterministic detection engines: YARA-X, Sigma, Lua, PowerShell, Python. "
        "Claude is an advisory escalation in Phase 2, never on the critical path."
    )
    add_image(doc, dp["threat_state_flow"], width_inches=5.5, caption="Figure 5.3 — ai-threat-state-graph: CAAD loop")
    add_body(doc,
        "Key insight: LLM is never on the critical detection path. This is the correct "
        "architecture for security-critical systems where latency and reliability matter more "
        "than AI flexibility."
    )

    # ai-procwatch-mcp
    set_heading(doc, "ai-procwatch-mcp — MCP Tool Registry + Two-Tier Triage", 2)
    add_body(doc,
        "Pattern: MCP server exposing 8 tools; behavioral genome tracked via ETW + eBPF; "
        "two-tier LLM classification (Ollama triage first → Claude escalation on high-confidence "
        "anomalies only)."
    )
    add_image(doc, dp["procwatch_flow"], width_inches=5.5, caption="Figure 5.4 — ai-procwatch-mcp: two-tier triage")
    add_body(doc,
        "Key insight: The two-tier triage pattern keeps most inference local (low cost, low "
        "latency) and escalates to a premium model only when the local model flags high confidence. "
        "Server-initiated sampling means escalation happens mid-capture without waiting for the user."
    )

    # ai-asset-sweeper
    set_heading(doc, "ai-asset-sweeper — Sandboxed Lua DSL (No LLM)", 2)
    add_body(doc,
        "Pattern: No LLM. Detection logic lives in Lua scripts; C++ orchestrator runs them in a "
        "sandboxed lua_State with whitelisted functions only. Hot-reloadable without recompiling "
        "the C++ host."
    )
    add_image(doc, dp["asset_sweeper_flow"], width_inches=5.5, caption="Figure 5.5 — ai-asset-sweeper: sandboxed Lua DSL")
    add_body(doc,
        "Key insight: The sandbox removes dangerous standard library functions — isolating the "
        "part that touches untrusted input (detector scripts) from the part that must not fail "
        "(the C++ host). This is the Error Kernel principle applied."
    )

    # ── Part 6: Quick Reference ───────────────────────────────────────────────
    set_heading(doc, "Part 6 — Quick Reference", 1)

    add_table(doc,
        ["Pattern", "Control Flow", "Memory", "Best For"],
        [
            ["REPL", "None", "None", "Simple chatbot"],
            ["ReAct", "LLM-driven loop", "In-context", "Tool use, search"],
            ["Plan-and-Execute", "Planner LLM", "State dict", "Long-horizon tasks"],
            ["Reflection", "Quality gate loop", "In-context", "Quality-sensitive generation"],
            ["Multi-Agent", "Orchestrator", "Shared state", "Parallelizable subtasks"],
            ["Tree of Thoughts", "BFS/DFS search", "In-context", "Optimization, proofs"],
            ["Memory-Augmented", "Any", "Vector DB / KB", "Long-term recall"],
            ["Event-Driven", "External event", "Persistent", "Background automation"],
            ["LangGraph Graph", "Code-defined", "Checkpointer", "Production workflows"],
            ["AutoGen GroupChat", "Emergent conv.", "Chat history", "Coding agents, research"],
            ["CrewAI Process", "Task pipeline", "Embedding store", "Role-based workflows"],
        ]
    )

    set_heading(doc, "The Key Architectural Principle", 2)
    add_body(doc,
        "LLM on the critical path = latency + reliability risk. Keep LLMs advisory or "
        "escalation-only for safety-critical, latency-sensitive, or regulated systems. Use "
        "deterministic engines as the primary layer; bring the LLM in only when confidence is "
        "low or human escalation is warranted."
    )
    add_body(doc,
        "This principle appears in ai-threat-state-graph (CAAD loop, LLM off critical path) "
        "and ai-procwatch-mcp (two-tier triage, Ollama first, Claude escalation only)."
    )

    # ── Save ──────────────────────────────────────────────────────────────────
    out_path = OUT_DIR / "ai_agent_patterns.docx"
    doc.save(str(out_path))
    print(f"\nDocument saved: {out_path}")
    return out_path


if __name__ == "__main__":
    diagram_paths = render_diagrams()
    build_doc(diagram_paths)
